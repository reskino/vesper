import os
import sys
import uuid
import json as _json
import base64
import logging
import shutil
import subprocess
import threading
import time
import tempfile
from flask import Flask, jsonify, request, send_file
from flask_cors import CORS

from config import AI_CONFIGS, FALLBACK_ORDER, get_active_model, set_active_model
from playwright_utils import (
    session_exists, get_session_info, delete_session, send_prompt, check_ai_available,
)
from history_manager import (
    add_message, get_messages, get_all_summaries, clear_messages, get_stats
)
from file_manager import (
    get_file_tree, read_file, write_file, create_file, delete_path,
    rename_path, get_language, LANGUAGE_MAP
)
from terminal_manager import exec_command, get_cwd, set_cwd, get_env_info
from agent import run_agent, get_status as get_agent_status, get_screenshot_path, stop_agent
from router import route as smart_route, explain as explain_route

logging.basicConfig(
    level=logging.DEBUG,
    format="%(asctime)s %(levelname)s %(name)s - %(message)s",
    handlers=[logging.StreamHandler(sys.stdout)]
)
logger = logging.getLogger(__name__)

app = Flask(__name__)
CORS(app, resources={r"/api/*": {"origins": "*"}})


@app.route("/api/healthz")
def health_check():
    return jsonify({"status": "ok"})


# ─── AI Proxy ────────────────────────────────────────────────────────────────

@app.route("/api/proxy/ais")
def list_ais():
    ais = []
    for ai_id, config in AI_CONFIGS.items():
        ais.append({
            "id": ai_id,
            "name": config["name"],
            "url": config["url"],
            "hasSession": session_exists(ai_id),
            "isAvailable": True,
            "icon": config.get("icon"),
            "models": config.get("models", []),
            "currentModel": get_active_model(ai_id),
            "authMode": config.get("auth_mode", "cookies"),
            "keyLabel": config.get("key_label"),
            "keyPrefix": config.get("key_prefix"),
            "keyUrl": config.get("key_url"),
            "keyUrlLabel": config.get("key_url_label"),
            "keyNote": config.get("key_note"),
        })
    return jsonify({"ais": ais})


@app.route("/api/proxy/set-model", methods=["POST"])
def set_model():
    data = request.get_json()
    ai_id = data.get("aiId")
    model_id = data.get("modelId")
    if not ai_id or not model_id:
        return jsonify({"error": "aiId and modelId are required"}), 400
    ok = set_active_model(ai_id, model_id)
    if not ok:
        return jsonify({"error": f"Invalid model '{model_id}' for AI '{ai_id}'"}), 400
    return jsonify({"success": True, "aiId": ai_id, "modelId": model_id})


@app.route("/api/proxy/route", methods=["POST"])
def route_prompt():
    """Preview which AI the router would choose for a given prompt (no message sent)."""
    data = request.get_json()
    prompt = data.get("prompt", "").strip()
    num_files = int(data.get("numFiles", 0))

    if not prompt:
        return jsonify({"error": "prompt is required"}), 400

    connected_ids = [aid for aid in AI_CONFIGS if session_exists(aid)]
    decision = smart_route(prompt, connected_ids, num_files=num_files)
    decision["explanation"] = explain_route(decision)
    return jsonify(decision)


@app.route("/api/proxy/ask", methods=["POST"])
def ask_ai():
    data = request.get_json()
    ai_id = data.get("aiId")
    prompt = data.get("prompt", "").strip()
    conversation_id = data.get("conversationId") or str(uuid.uuid4())
    use_fallback = data.get("fallback", True)

    if not ai_id or not prompt:
        return jsonify({"error": "aiId and prompt are required"}), 400

    # ── Smart routing when aiId is "__auto__" ─────────────────────────────────
    routing_decision = None
    if ai_id == "__auto__":
        connected_ids = [aid for aid in AI_CONFIGS if session_exists(aid)]
        routing_decision = smart_route(prompt, connected_ids)
        logger.info(
            "Smart router → %s (confidence %.2f) — %s",
            routing_decision["aiId"], routing_decision["confidence"], routing_decision["reason"]
        )
        ai_id = routing_decision["aiId"]
    # ─────────────────────────────────────────────────────────────────────────

    add_message(ai_id, "user", prompt, conversation_id)

    start_time = time.time()
    fallback_used = False
    tried_ais = []
    ai_errors: dict = {}

    ais_to_try = [ai_id]
    if use_fallback:
        for fallback_ai in FALLBACK_ORDER:
            if fallback_ai != ai_id and fallback_ai not in ais_to_try:
                ais_to_try.append(fallback_ai)

    for current_ai in ais_to_try:
        if not session_exists(current_ai):
            tried_ais.append(current_ai)
            ai_errors[current_ai] = "no session"
            continue

        if current_ai != ai_id:
            fallback_used = True
            logger.info("Falling back to %s (tried: %s)", current_ai, tried_ais)

        success, response_text, error = send_prompt(current_ai, prompt)
        elapsed_ms = int((time.time() - start_time) * 1000)

        if success and response_text:
            add_message(ai_id, "assistant", response_text, conversation_id)
            return jsonify({
                "success": True,
                "aiId": current_ai,
                "response": response_text,
                "conversationId": conversation_id,
                "elapsedMs": elapsed_ms,
                "fallbackUsed": fallback_used,
                "routingDecision": routing_decision,
                "error": None,
            })

        tried_ais.append(current_ai)
        ai_errors[current_ai] = error
        logger.warning("AI %s failed: %s", current_ai, error)

    elapsed_ms = int((time.time() - start_time) * 1000)
    has_any_session = any(session_exists(a) for a in ais_to_try)
    if not has_any_session:
        if not use_fallback:
            cfg = AI_CONFIGS.get(ai_id, {})
            ai_name = cfg.get("name", ai_id)
            auth_mode = cfg.get("auth_mode", "")
            if auth_mode in ("api_key", "api_key_or_cookies"):
                err_msg = f"{ai_name} is not connected. Go to Sessions and add your API key for {ai_name}."
            else:
                err_msg = f"{ai_name} is not connected. Go to Sessions and import your browser cookies."
        else:
            err_msg = "No active sessions found. Go to the Sessions page and connect at least one AI provider."
    else:
        details = "; ".join(
            f"{ai}: {msg}" for ai, msg in ai_errors.items() if msg != "no session"
        )
        err_msg = f"All AIs failed. {details}" if details else "All AIs failed — session cookies may have expired."
    return jsonify({
        "success": False,
        "aiId": ai_id,
        "response": "",
        "conversationId": conversation_id,
        "elapsedMs": elapsed_ms,
        "fallbackUsed": fallback_used,
        "error": err_msg,
    })


@app.route("/api/proxy/ask-with-context", methods=["POST"])
def ask_ai_with_context():
    data = request.get_json()
    ai_id = data.get("aiId")
    user_prompt = data.get("prompt", "").strip()
    files = data.get("files", [])
    action = data.get("action")
    agent_type = data.get("agentType")          # e.g. "code_surgeon", "scholar" …
    conversation_id = data.get("conversationId") or str(uuid.uuid4())
    use_fallback = data.get("fallback", True)

    if not ai_id or not user_prompt:
        return jsonify({"error": "aiId and prompt are required"}), 400

    # ── Smart routing when aiId is "__auto__" ─────────────────────────────────
    routing_decision_ctx = None
    if ai_id == "__auto__":
        connected_ids = [aid for aid in AI_CONFIGS if session_exists(aid)]
        routing_decision_ctx = smart_route(user_prompt, connected_ids, num_files=len(files))
        logger.info(
            "Smart router (ctx) → %s — %s",
            routing_decision_ctx["aiId"], routing_decision_ctx["reason"]
        )
        ai_id = routing_decision_ctx["aiId"]
    # ─────────────────────────────────────────────────────────────────────────

    # ── Action-level prefixes (appended before the user message) ─────────────
    ACTION_PREFIXES = {
        "fix":      "Please analyze the following code and fix any bugs. Explain what you changed.\n\n",
        "explain":  "Please explain the following code in detail.\n\n",
        "test":     "Please write comprehensive unit tests for the following code.\n\n",
        "refactor": "Please refactor the following code to improve readability and maintainability. "
                    "Preserve behaviour; only improve structure, naming, and clarity.\n\n",
        "suggest":  "Please review the following code and suggest improvements.\n\n",
        "debug":    "Please help debug the following code and provide a fix.\n\n",
        "document": "Please add comprehensive documentation and docstrings to the following code. "
                    "Use the appropriate format for the language (JSDoc, Google-style Python, etc.).\n\n",
    }

    # ── Agent-type system prompts (prepended before action prefix) ────────────
    # These tell the AI which specialist persona to embody, giving richer
    # persona context than the short banner already prepended by the frontend.
    AGENT_SYSTEM_PROMPTS = {
        "code_surgeon": (
            "You are Vesper Code Surgeon — a senior software engineer specialising in "
            "precise, surgical code improvements. Your priorities: correctness first, "
            "readability second, performance third. Always explain what you changed and why. "
            "Prefer minimal diffs; never rewrite working code unnecessarily.\n\n"
        ),
        "scholar": (
            "You are Vesper Research Scholar — a rigorous academic assistant. "
            "Provide well-structured, citation-backed answers. Use numbered sections, "
            "bullet points, and code examples where relevant. Acknowledge uncertainty "
            "explicitly and distinguish facts from opinions.\n\n"
        ),
        "search_master": (
            "You are Vesper Search Master — a deep research assistant. "
            "Surface the most authoritative, up-to-date sources. "
            "Always include real URLs or package references where possible. "
            "Summarise findings concisely then provide detail on request.\n\n"
        ),
        "docs_weaver": (
            "You are Vesper Docs Weaver — a technical writing specialist. "
            "Produce clear, beautifully structured documentation. "
            "Use Markdown headings, tables, and code blocks. "
            "Tailor the tone to the audience (developer docs vs end-user guides).\n\n"
        ),
        "orchestrator": (
            "You are Vesper Orchestrator — a full-stack architect with broad expertise. "
            "Break large problems into concrete, ordered steps. "
            "Consider scalability, security, and maintainability in every design decision. "
            "Produce production-quality, complete implementations — never stubs.\n\n"
        ),
    }

    agent_prefix = AGENT_SYSTEM_PROMPTS.get(agent_type, "") if agent_type else ""
    action_prefix = ACTION_PREFIXES.get(action, "") if action else ""
    prefix = agent_prefix + action_prefix

    # ── Build file context ────────────────────────────────────────────────────
    file_context = ""
    if files:
        # Separate the special workspace context file from regular attachments
        ws_context_parts = []
        regular_files = []
        for f in files:
            if f.get("path") in ("__workspace_context__", "__imported_project_context__"):
                ws_context_parts.append(f)
            else:
                regular_files.append(f)

        if ws_context_parts:
            file_context += "\n\n--- Workspace Context ---\n"
            for f in ws_context_parts:
                file_context += f"\n{f['content']}\n"

        if regular_files:
            file_context += "\n\n--- Attached Files ---\n"
            for f in regular_files:
                lang = f.get("language") or get_language(f.get("path", ""))
                file_context += f"\n**File: `{f['path']}`**\n```{lang}\n{f['content']}\n```\n"

    full_prompt = prefix + user_prompt + file_context
    add_message(ai_id, "user", full_prompt, conversation_id)

    start_time = time.time()
    fallback_used = False
    tried_ais = []

    ais_to_try = [ai_id]
    if use_fallback:
        for fallback_ai in FALLBACK_ORDER:
            if fallback_ai != ai_id and fallback_ai not in ais_to_try:
                ais_to_try.append(fallback_ai)

    ai_errors: dict = {}

    for current_ai in ais_to_try:
        if not session_exists(current_ai):
            tried_ais.append(current_ai)
            ai_errors[current_ai] = "no session"
            continue

        if current_ai != ai_id:
            fallback_used = True

        success, response_text, error = send_prompt(current_ai, full_prompt)
        elapsed_ms = int((time.time() - start_time) * 1000)

        if success and response_text:
            add_message(ai_id, "assistant", response_text, conversation_id)
            return jsonify({
                "success": True,
                "aiId": current_ai,
                "response": response_text,
                "conversationId": conversation_id,
                "elapsedMs": elapsed_ms,
                "fallbackUsed": fallback_used,
                "routingDecision": routing_decision_ctx,
                "error": None,
            })

        tried_ais.append(current_ai)
        ai_errors[current_ai] = error
        logger.warning("AI %s failed: %s", current_ai, error)

    elapsed_ms = int((time.time() - start_time) * 1000)
    has_any_session = any(session_exists(a) for a in ais_to_try)
    if not has_any_session:
        if not use_fallback:
            cfg = AI_CONFIGS.get(ai_id, {})
            ai_name = cfg.get("name", ai_id)
            auth_mode = cfg.get("auth_mode", "")
            if auth_mode in ("api_key", "api_key_or_cookies"):
                err_msg = f"{ai_name} is not connected. Go to Sessions and add your API key for {ai_name}."
            else:
                err_msg = f"{ai_name} is not connected. Go to Sessions and import your browser cookies."
        else:
            err_msg = "No active sessions found. Go to the Sessions page and connect at least one AI provider."
    else:
        details = "; ".join(
            f"{ai}: {msg}" for ai, msg in ai_errors.items() if msg != "no session"
        )
        err_msg = f"All AIs failed. {details}" if details else "All AIs failed — session cookies may have expired."
    return jsonify({
        "success": False,
        "aiId": ai_id,
        "response": "",
        "conversationId": conversation_id,
        "elapsedMs": elapsed_ms,
        "fallbackUsed": fallback_used,
        "error": err_msg,
    })


@app.route("/api/proxy/execute", methods=["POST"])
def execute_code():
    data = request.get_json()
    code = data.get("code", "").strip()
    language = data.get("language", "python")
    timeout = min(int(data.get("timeout", 30)), 60)

    if not code:
        return jsonify({"error": "code is required"}), 400

    start_time = time.time()
    allowed_languages = {
        "python": ["python3", "-c"],
        "javascript": ["node", "-e"],
        "bash": ["bash", "-c"]
    }
    if language not in allowed_languages:
        return jsonify({"error": f"Unsupported language: {language}"}), 400

    cmd = allowed_languages[language] + [code]
    try:
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=timeout, cwd="/tmp")
        elapsed_ms = int((time.time() - start_time) * 1000)
        return jsonify({
            "success": result.returncode == 0,
            "stdout": result.stdout,
            "stderr": result.stderr,
            "exitCode": result.returncode,
            "elapsedMs": elapsed_ms,
        })
    except subprocess.TimeoutExpired:
        return jsonify({"success": False, "stdout": "", "stderr": f"Timed out after {timeout}s", "exitCode": 1, "elapsedMs": timeout * 1000})
    except Exception as e:
        return jsonify({"success": False, "stdout": "", "stderr": str(e), "exitCode": 1, "elapsedMs": 0})


# ─── Sessions ────────────────────────────────────────────────────────────────

@app.route("/api/sessions")
def list_sessions():
    sessions = []
    for ai_id, config in AI_CONFIGS.items():
        info = get_session_info(ai_id)
        sessions.append({"aiId": ai_id, "aiName": config["name"], **info})
    return jsonify({"sessions": sessions})


# ─── Browser session worker management ───────────────────────────────────────
_browser_workers: dict = {}
_workers_lock = threading.Lock()

XVFB_RUN = shutil.which("xvfb-run")
WORKER_SCRIPT = os.path.join(os.path.dirname(__file__), "session_browser_worker.py")


def _launch_browser_worker(ai_id: str) -> tuple:
    work_dir = os.path.join(tempfile.gettempdir(), f"vesper_browser_{ai_id}")
    os.makedirs(work_dir, exist_ok=True)
    cmd_path = os.path.join(work_dir, "command.txt")
    if os.path.exists(cmd_path):
        os.remove(cmd_path)

    log_path = os.path.join(work_dir, "worker.log")
    log_file = open(log_path, "w")
    cmd = [sys.executable, WORKER_SCRIPT, ai_id, work_dir]
    proc = subprocess.Popen(cmd, stdout=log_file, stderr=log_file)
    with _workers_lock:
        _browser_workers[ai_id] = {"proc": proc, "work_dir": work_dir}
    return proc, work_dir


def _send_worker_command(ai_id: str, payload: dict) -> bool:
    with _workers_lock:
        info = _browser_workers.get(ai_id)
    if not info:
        return False
    cmd_path = os.path.join(info["work_dir"], "command.txt")
    with open(cmd_path, "w") as f:
        f.write(_json.dumps(payload))
    return True


@app.route("/api/sessions/create", methods=["POST"])
def create_session():
    data = request.get_json()
    ai_id = data.get("aiId")
    if not ai_id:
        return jsonify({"success": False, "message": "aiId is required"}), 400
    if ai_id not in AI_CONFIGS:
        return jsonify({"success": False, "message": f"Unknown AI: {ai_id}"}), 400

    with _workers_lock:
        old = _browser_workers.pop(ai_id, None)
    if old:
        try:
            old["proc"].terminate()
        except Exception:
            pass

    try:
        _launch_browser_worker(ai_id)
    except Exception as e:
        logger.error("Failed to launch browser worker: %s", e)
        return jsonify({"success": False, "message": str(e)}), 500

    return jsonify({
        "success": True,
        "message": f"Browser launched for {AI_CONFIGS[ai_id]['name']}. Log in and click Save.",
        "aiId": ai_id,
    })


@app.route("/api/sessions/browser-status/<ai_id>", methods=["GET"])
def browser_status(ai_id):
    with _workers_lock:
        info = _browser_workers.get(ai_id)
    if not info:
        return jsonify({"active": False, "status": "idle"})

    proc = info["proc"]
    if proc.poll() is not None:
        with _workers_lock:
            _browser_workers.pop(ai_id, None)

    status_file = os.path.join(info["work_dir"], "status.json")
    if os.path.exists(status_file):
        try:
            with open(status_file) as f:
                status = _json.load(f)
            return jsonify({"active": proc.poll() is None, **status})
        except Exception:
            pass
    return jsonify({"active": proc.poll() is None, "status": "starting"})


@app.route("/api/sessions/browser-screenshot/<ai_id>", methods=["GET"])
def browser_screenshot(ai_id):
    """Return the latest browser screenshot as a raw PNG (no base64 overhead)."""
    # Try the active worker first
    with _workers_lock:
        info = _browser_workers.get(ai_id)

    # Fall back to the last known work dir even if worker has exited
    if not info:
        work_dir = os.path.join(tempfile.gettempdir(), f"vesper_browser_{ai_id}")
        screenshot_file = os.path.join(work_dir, "latest.png")
    else:
        screenshot_file = os.path.join(info["work_dir"], "latest.png")

    if not os.path.exists(screenshot_file):
        return ("", 204)  # No content yet — browser will keep polling

    try:
        response = send_file(screenshot_file, mimetype="image/png")
        response.headers["Cache-Control"] = "no-store, no-cache, must-revalidate"
        response.headers["Pragma"] = "no-cache"
        return response
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/sessions/browser-action/<ai_id>", methods=["POST"])
def browser_action(ai_id):
    payload = request.get_json() or {}
    ok = _send_worker_command(ai_id, payload)
    if not ok:
        return jsonify({"success": False, "message": "No active browser session"}), 404
    return jsonify({"success": True})


@app.route("/api/sessions/<ai_id>/delete", methods=["DELETE"])
def delete_session_route(ai_id):
    success, message = delete_session(ai_id)
    return jsonify({"success": success, "message": message, "aiId": ai_id})


@app.route("/api/sessions/import", methods=["POST"])
def import_session():
    from config import SESSIONS_DIR
    data = request.get_json()
    ai_id = data.get("aiId")
    state_json = data.get("stateJson")

    if not ai_id or not state_json:
        return jsonify({"success": False, "message": "aiId and stateJson are required"}), 400
    if ai_id not in AI_CONFIGS:
        return jsonify({"success": False, "message": f"Unknown AI: {ai_id}"}), 400

    try:
        parsed = _json.loads(state_json) if isinstance(state_json, str) else state_json
    except Exception as exc:
        return jsonify({"success": False, "message": f"Invalid JSON: {exc}"}), 400

    # ── Normalise to Playwright storage_state format ──────────────────────────
    # Cookie Editor exports a flat array: [{name, value, domain, path, ...}, ...]
    # Playwright expects:                 {"cookies": [...], "origins": [...]}
    if isinstance(parsed, list):
        # Convert Cookie Editor / Netscape-JSON array → Playwright format
        # Cookie Editor / browser export values → Playwright's required enum
        _SAME_SITE_MAP = {
            "strict":         "Strict",
            "lax":            "Lax",
            "none":           "None",
            "no_restriction": "None",   # Cookie Editor calls SameSite=None "no_restriction"
            "unspecified":    "Lax",
            "":               "Lax",
        }

        def _norm_cookie(c: dict) -> dict:
            """Map Cookie Editor fields to Playwright cookie fields."""
            raw_ss = str(c.get("sameSite") or "").strip().lower()
            out: dict = {
                "name":     c.get("name", ""),
                "value":    c.get("value", ""),
                "domain":   c.get("domain", ""),
                "path":     c.get("path", "/"),
                "secure":   bool(c.get("secure", False)),
                "httpOnly": bool(c.get("httpOnly", False)),
                "sameSite": _SAME_SITE_MAP.get(raw_ss, "Lax"),
            }
            # Playwright uses "expires" (float epoch); Cookie Editor uses "expirationDate"
            exp = c.get("expirationDate") or c.get("expires")
            if exp is not None:
                out["expires"] = float(exp)
            else:
                out["expires"] = -1
            return out

        state = {"cookies": [_norm_cookie(c) for c in parsed if isinstance(c, dict)], "origins": []}
    elif isinstance(parsed, dict):
        state = parsed
        if "cookies" not in state and "origins" not in state:
            return jsonify({
                "success": False,
                "message": "JSON must contain a 'cookies' array or be a Cookie Editor export (flat array).",
            }), 400
    else:
        return jsonify({"success": False, "message": "JSON must be an object or a Cookie Editor cookie array."}), 400

    os.makedirs(SESSIONS_DIR, exist_ok=True)
    session_path = os.path.join(SESSIONS_DIR, f"{ai_id}_state.json")
    try:
        with open(session_path, "w") as f:
            _json.dump(state, f, indent=2)
        # ── Also persist to KV so sessions survive filesystem resets ──────────
        from key_store import save_session_state as _ks_save_session
        _ks_save_session(ai_id, state)
        n = len(state.get("cookies", []))
        return jsonify({
            "success": True,
            "message": f"Imported {n} cookie(s) for {AI_CONFIGS[ai_id]['name']}.",
            "aiId": ai_id,
        })
    except Exception as exc:
        return jsonify({"success": False, "message": str(exc)}), 500


@app.route("/api/sessions/import-key", methods=["POST"])
def import_api_key():
    """Store an API key for any AI provider that supports API key auth."""
    from key_store import save_api_key
    data = request.get_json()
    ai_id  = data.get("aiId")
    api_key = (data.get("apiKey") or "").strip()

    if not ai_id or not api_key:
        return jsonify({"success": False, "message": "aiId and apiKey are required"}), 400
    if ai_id not in AI_CONFIGS:
        return jsonify({"success": False, "message": f"Unknown AI: {ai_id}"}), 400

    cfg = AI_CONFIGS[ai_id]
    if cfg.get("auth_mode") not in ("api_key", "api_key_or_cookies"):
        return jsonify({"success": False, "message": f"{cfg['name']} does not support API key authentication."}), 400

    ok = save_api_key(ai_id, api_key)
    if ok:
        return jsonify({"success": True, "message": f"API key saved for {cfg['name']}.", "aiId": ai_id})
    return jsonify({"success": False, "message": "Failed to save API key."}), 500


def get_api_key(ai_id: str) -> str:
    """Return the stored API key for an AI, or '' if none."""
    from key_store import load_api_key
    return load_api_key(ai_id)


@app.route("/api/sessions/verify/<ai_id>", methods=["GET"])
def verify_session(ai_id):
    """Verify saved credentials are valid and return the logged-in username."""
    if ai_id not in AI_CONFIGS:
        return jsonify({"success": False, "error": f"Unknown AI: {ai_id}"}), 400

    from playwright_utils import session_exists, get_session_path

    cfg = AI_CONFIGS[ai_id]

    # No-auth providers (Pollinations, LLM7, etc.) — always active, no key needed
    if cfg.get("auth_mode") == "none":
        return jsonify({"success": True, "username": "No login required", "authMode": "none"})

    # ── API key providers ─────────────────────────────────────────────────────
    # Map of ai_id → (models_url, headers_fn) for live ping verification.
    # Use a Bearer token pattern for all OpenAI-compat providers.
    def _bearer(key): return {"Authorization": f"Bearer {key}"}
    _VERIFY_PINGS = {
        "chatgpt":     lambda k: ("https://api.openai.com/v1/models",                        _bearer(k)),
        "claude":      lambda k: ("https://api.anthropic.com/v1/models",                     {"x-api-key": k, "anthropic-version": "2023-06-01"}),
        "groq":        lambda k: ("https://api.groq.com/openai/v1/models",                   {**_bearer(k), "User-Agent": "Vesper"}),
        "gemini":      lambda k: (f"https://generativelanguage.googleapis.com/v1beta/models?key={k}", {}),
        "openrouter":  lambda k: ("https://openrouter.ai/api/v1/models",                     _bearer(k)),
        "mistral":     lambda k: ("https://api.mistral.ai/v1/models",                        _bearer(k)),
        "cerebras":    lambda k: ("https://api.cerebras.ai/v1/models",                       _bearer(k)),
        "deepseek":    lambda k: ("https://api.deepseek.com/v1/models",                      _bearer(k)),
        "cohere":      lambda k: ("https://api.cohere.com/v1/models",                        _bearer(k)),
        "nvidia":      lambda k: ("https://integrate.api.nvidia.com/v1/models",              _bearer(k)),
        "github":      lambda k: ("https://models.inference.ai.azure.com/v1/models",         _bearer(k)),
        "huggingface": lambda k: ("https://router.huggingface.co/v1/models",                 _bearer(k)),
        "kluster":     lambda k: ("https://api.kluster.ai/v1/models",                        _bearer(k)),
        "siliconflow": lambda k: ("https://api.siliconflow.cn/v1/models",                    _bearer(k)),
        "zhipu":       lambda k: ("https://open.bigmodel.cn/api/paas/v4/models",             _bearer(k)),
    }

    if cfg.get("auth_mode") in ("api_key", "api_key_or_cookies"):
        api_key = get_api_key(ai_id)
        if api_key:
            import urllib.request, urllib.error as _ue
            ping_fn = _VERIFY_PINGS.get(ai_id)
            if ping_fn:
                try:
                    url, headers = ping_fn(api_key)
                    req = urllib.request.Request(url, headers=headers)
                    with urllib.request.urlopen(req, timeout=10) as r:
                        _json.loads(r.read())
                    return jsonify({"success": True, "username": "API Key ✓", "authMode": "api_key"})
                except _ue.HTTPError as exc:
                    if exc.code in (401, 403):
                        return jsonify({"success": False, "error": "API key invalid or expired"})
                    # 429, 5xx, etc. → key exists, provider just busy
                    return jsonify({"success": True, "username": "API Key ✓", "authMode": "api_key"})
                except Exception as exc:
                    # Network timeout or other transient error — key IS stored, treat as valid
                    logger.warning("Verify ping for %s failed (non-auth): %s", ai_id, exc)
                    return jsonify({"success": True, "username": "API Key ✓", "authMode": "api_key"})
            else:
                # Provider has no ping URL — key existence is sufficient
                return jsonify({"success": True, "username": "API Key ✓", "authMode": "api_key"})

    if not session_exists(ai_id):
        return jsonify({"success": False, "error": "No API key or session saved"})

    session_path = get_session_path(ai_id)

    try:
        from web_session_client import _load_cookies, _impersonate_session
        cookies = _load_cookies(session_path, "")
        if not cookies:
            return jsonify({"success": False, "error": "Session file has no cookies"})

        sess = _impersonate_session(cookies)

        if ai_id == "chatgpt":
            # Use backend-anon sentinel instead of /api/auth/session (Cloudflare blocks the latter)
            device_id = str(uuid.uuid4())
            sess.headers.update({
                "oai-device-id": device_id,
                "Origin": "https://chatgpt.com",
                "Referer": "https://chatgpt.com/",
                "sec-fetch-site": "same-origin",
            })
            req_resp = sess.post(
                "https://chatgpt.com/backend-anon/sentinel/chat-requirements",
                json={}, timeout=15,
            )
            if req_resp.status_code == 200:
                # Cookies are valid enough to reach the sentinel endpoint
                # Try to get a display name from /api/me or user-profile endpoints
                name = "Connected"
                for profile_url in [
                    "https://chatgpt.com/backend-api/me",
                    "https://chatgpt.com/backend-anon/me",
                ]:
                    try:
                        me = sess.get(profile_url, timeout=10)
                        if me.status_code == 200:
                            d = me.json()
                            n = d.get("name") or d.get("email") or d.get("username")
                            if n:
                                name = n
                                break
                    except Exception:
                        continue
                return jsonify({"success": True, "username": name})
            return jsonify({"success": False,
                            "error": f"ChatGPT session invalid ({req_resp.status_code}) — please re-import cookies"})

        elif ai_id == "claude":
            resp = sess.get("https://claude.ai/api/organizations", timeout=15)
            if resp.status_code == 200:
                orgs = resp.json()
                if isinstance(orgs, list) and orgs:
                    # Get account info for the actual email
                    acc = sess.get("https://claude.ai/api/account", timeout=10)
                    email = ""
                    if acc.status_code == 200:
                        email = acc.json().get("email", "")
                    name = email or orgs[0].get("name", "Connected")
                    return jsonify({"success": True, "username": name, "email": email})
            return jsonify({"success": False,
                            "error": f"Session invalid ({resp.status_code}) — please re-import cookies"})

        elif ai_id == "grok":
            # Grok is behind Cloudflare — curl_cffi GET requests are challenged.
            # Use a real Playwright browser (cookies already loaded) to hit the
            # /api/me endpoint from inside the page context, which bypasses CF.
            grok_username = None
            try:
                from playwright.sync_api import sync_playwright as _spw
                _grok_ua = (
                    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
                    "AppleWebKit/537.36 (KHTML, like Gecko) "
                    "Chrome/136.0.0.0 Safari/537.36"
                )
                with _spw() as _pw:
                    from config import find_chromium as _fc  # noqa: PLC0415
                    _chrome_exe = _fc()
                    _lkw: dict = {
                        "headless": True,
                        "args": ["--no-sandbox", "--disable-setuid-sandbox",
                                 "--disable-dev-shm-usage", "--disable-gpu",
                                 "--disable-blink-features=AutomationControlled"],
                    }
                    if _chrome_exe:
                        _lkw["executable_path"] = _chrome_exe
                    _br = _pw.chromium.launch(**_lkw)
                    _ctx = _br.new_context(
                        viewport={"width": 1280, "height": 900},
                        user_agent=_grok_ua,
                        storage_state=session_path,
                    )
                    _ctx.add_init_script(
                        "Object.defineProperty(navigator,'webdriver',{get:()=>undefined});"
                    )
                    _pg = _ctx.new_page()
                    _pg.set_default_timeout(30_000)
                    try:
                        _pg.goto("https://grok.com/", timeout=30_000,
                                 wait_until="domcontentloaded")
                        _pg.wait_for_timeout(2_000)
                    except Exception:
                        pass
                    # Grok doesn't expose a simple /api/me JSON endpoint.
                    # Extract username directly from the page DOM.
                    if not grok_username:
                        try:
                            _handle = _pg.evaluate(
                                """() => {
                                    // Walk all visible text nodes looking for @handle patterns
                                    // or find elements with username-like aria-labels / data attributes
                                    const selectors = [
                                        '[data-testid*="user"] [aria-label]',
                                        '[class*="screen"] span',
                                        '[aria-label*="@"]',
                                        'a[href*="/profile"] span',
                                        'button[aria-label*="account"] span',
                                        'nav a[aria-label] span',
                                    ];
                                    for (const sel of selectors) {
                                        const el = document.querySelector(sel);
                                        if (el) {
                                            const t = el.textContent.trim();
                                            if (t.length > 0) return t;
                                        }
                                    }
                                    // Look for any @mention pattern in visible text
                                    const walker = document.createTreeWalker(
                                        document.body, NodeFilter.SHOW_TEXT, null);
                                    let node;
                                    while (node = walker.nextNode()) {
                                        const m = node.textContent.match(/@([A-Za-z0-9_]{1,50})/);
                                        if (m) return '@' + m[1];
                                    }
                                    return null;
                                }"""
                            )
                            if _handle and len(_handle) > 0:
                                grok_username = _handle
                        except Exception:
                            pass
                    # If DOM extraction failed, fall back to x-userid cookie short ID
                    if not grok_username:
                        for _ck in _ctx.cookies():
                            if _ck.get("name") == "x-userid" and _ck.get("value"):
                                grok_username = f"SuperGrok ({_ck['value'][:8]}…)"
                                break
                    _br.close()
            except Exception as _pw_err:
                logger.warning("Grok Playwright verify failed: %s", _pw_err)

            if grok_username:
                # Prefix @ if it looks like a Twitter handle and doesn't already have it
                if grok_username and not grok_username.startswith("@") and "@" not in grok_username and " " not in grok_username:
                    grok_username = f"@{grok_username}"
                return jsonify({"success": True, "username": grok_username})
            # Cookies exist and we loaded the page — mark as connected
            return jsonify({"success": True, "username": "Connected",
                            "warning": "Could not fetch Grok username"})

        return jsonify({"success": False, "error": f"No verifier for {ai_id}"})

    except Exception as exc:
        logger.error("verify_session error: %s", exc, exc_info=True)
        return jsonify({"success": False, "error": str(exc)})


# ─── Model Validation ────────────────────────────────────────────────────────

@app.route("/api/proxy/validate-models", methods=["POST"])
def validate_models():
    """
    For each stored API key, call the provider's live /models endpoint and
    compare against the model IDs declared in config.py.

    Returns per-provider: valid (in both), stale (in config but not live),
    live_only (in live but not config, i.e. new models available).
    """
    import urllib.request as _req_lib
    import urllib.error as _ue

    _PROVIDERS = {
        "chatgpt":     ("openai_compat", "https://api.openai.com/v1/models",                        {"Authorization": "Bearer {key}"}),
        "groq":        ("openai_compat", "https://api.groq.com/openai/v1/models",                   {"Authorization": "Bearer {key}", "User-Agent": "Vesper"}),
        "gemini":      ("gemini",        "https://generativelanguage.googleapis.com/v1beta/models?key={key}&pageSize=200", {}),
        "claude":      ("openai_compat", "https://api.anthropic.com/v1/models",                     {"x-api-key": "{key}", "anthropic-version": "2023-06-01"}),
        "openrouter":  ("openai_compat", "https://openrouter.ai/api/v1/models",                     {"Authorization": "Bearer {key}"}),
        "mistral":     ("openai_compat", "https://api.mistral.ai/v1/models",                        {"Authorization": "Bearer {key}"}),
        "cerebras":    ("openai_compat", "https://api.cerebras.ai/v1/models",                       {"Authorization": "Bearer {key}"}),
        "together":    ("together",      "https://api.together.xyz/v1/models",                       {"Authorization": "Bearer {key}"}),
        "deepseek":    ("openai_compat", "https://api.deepseek.com/v1/models",                      {"Authorization": "Bearer {key}"}),
        "cohere":      ("cohere",        "https://api.cohere.com/v1/models",                        {"Authorization": "Bearer {key}"}),
        "nvidia":      ("openai_compat", "https://integrate.api.nvidia.com/v1/models",              {"Authorization": "Bearer {key}"}),
        "github":      ("openai_compat", "https://models.inference.ai.azure.com/v1/models",         {"Authorization": "Bearer {key}"}),
        "huggingface": ("openai_compat", "https://router.huggingface.co/v1/models",                 {"Authorization": "Bearer {key}"}),
        "kluster":     ("openai_compat", "https://api.kluster.ai/v1/models",                        {"Authorization": "Bearer {key}"}),
        "siliconflow": ("openai_compat", "https://api.siliconflow.cn/v1/models",                    {"Authorization": "Bearer {key}"}),
        "zhipu":       ("openai_compat", "https://open.bigmodel.cn/api/paas/v4/models",             {"Authorization": "Bearer {key}"}),
    }

    results = {}

    # Handle no-auth providers (e.g. llm7, pollinations) — always report as live
    for ai_id, cfg in AI_CONFIGS.items():
        if cfg.get("auth_mode") == "none":
            config_ids = sorted(m["id"] for m in cfg.get("models", []) if m["id"] != "__auto__")
            results[ai_id] = {"status": "ok", "valid": config_ids, "stale": [], "live_only": [], "note": "no_auth"}

    for ai_id, (fmt, url_tpl, hdr_tpl) in _PROVIDERS.items():
        cfg = AI_CONFIGS.get(ai_id)
        if not cfg:
            continue
        auth_mode = cfg.get("auth_mode", "api_key")
        api_key = get_api_key(ai_id) if auth_mode != "none" else ""
        if auth_mode != "none" and not api_key:
            results[ai_id] = {"status": "no_key"}
            continue

        config_ids = {m["id"] for m in cfg.get("models", []) if m["id"] != "__auto__"}
        url = url_tpl.replace("{key}", api_key or "")
        headers = {k: v.replace("{key}", api_key or "") for k, v in hdr_tpl.items()}

        try:
            req = _req_lib.Request(url, headers=headers)
            with _req_lib.urlopen(req, timeout=12) as resp:
                data = _json.loads(resp.read())

            # Parse live model IDs depending on response shape
            live_ids: set[str] = set()
            if fmt == "openai_compat":
                if isinstance(data, list):
                    live_ids = {m["id"] for m in data if isinstance(m, dict)}
                else:
                    live_ids = {m["id"] for m in data.get("data", []) if isinstance(m, dict)}
            elif fmt == "gemini":
                live_ids = {m["name"].split("/")[-1] for m in data.get("models", []) if isinstance(m, dict) and "name" in m}
            elif fmt == "together":
                if isinstance(data, list):
                    live_ids = {m["id"] for m in data if isinstance(m, dict)}
                else:
                    live_ids = {m["id"] for m in data.get("data", []) if isinstance(m, dict)}
            elif fmt == "cohere":
                live_ids = {m.get("name", m.get("id", "")) for m in data.get("models", []) if isinstance(m, dict)}

            stale    = sorted(config_ids - live_ids)
            live_only = sorted(live_ids - config_ids)
            valid    = sorted(config_ids & live_ids)

            results[ai_id] = {
                "status":    "ok",
                "valid":     valid,
                "stale":     stale,
                "live_only": live_only[:20],  # cap to keep payload small
            }

        except _ue.HTTPError as exc:
            body = ""
            try: body = exc.read().decode(errors="replace")[:120]
            except Exception: pass
            results[ai_id] = {"status": "error", "error": f"HTTP {exc.code}: {body}"}
        except Exception as exc:
            results[ai_id] = {"status": "error", "error": str(exc)[:200]}

    return jsonify({"results": results})


# ─── History ─────────────────────────────────────────────────────────────────

@app.route("/api/history")
def list_history():
    summaries = get_all_summaries(AI_CONFIGS)
    return jsonify({"conversations": summaries})


@app.route("/api/history/stats")
def history_stats():
    stats = get_stats(AI_CONFIGS)
    return jsonify(stats)


@app.route("/api/history/<ai_id>")
def get_history(ai_id):
    messages = get_messages(ai_id)
    return jsonify({"aiId": ai_id, "messages": messages})


@app.route("/api/history/<ai_id>", methods=["DELETE"])
def clear_history_route(ai_id):
    clear_messages(ai_id)
    return jsonify({"success": True, "message": f"History cleared for {ai_id}", "aiId": ai_id})


# ─── File System ─────────────────────────────────────────────────────────────

@app.route("/api/files/tree")
def file_tree():
    path = request.args.get("path", ".")
    depth = min(int(request.args.get("depth", 3)), 6)
    try:
        return jsonify(get_file_tree(path, depth))
    except Exception as e:
        return jsonify({"error": str(e)}), 400


@app.route("/api/files/read")
def file_read():
    path = request.args.get("path")
    if not path:
        return jsonify({"error": "path is required"}), 400
    try:
        return jsonify(read_file(path))
    except FileNotFoundError as e:
        return jsonify({"error": str(e)}), 404
    except Exception as e:
        return jsonify({"error": str(e)}), 400


@app.route("/api/files/write", methods=["POST"])
def file_write():
    data = request.get_json()
    path = data.get("path")
    if not path:
        return jsonify({"error": "path is required"}), 400
    try:
        return jsonify(write_file(path, data.get("content", "")))
    except Exception as e:
        return jsonify({"error": str(e)}), 400


@app.route("/api/files/create", methods=["POST"])
def file_create():
    data = request.get_json()
    path = data.get("path")
    if not path:
        return jsonify({"error": "path is required"}), 400
    try:
        return jsonify(create_file(path, data.get("type", "file"), data.get("content")))
    except Exception as e:
        return jsonify({"error": str(e)}), 400


@app.route("/api/files/delete", methods=["DELETE"])
def file_delete():
    path = request.args.get("path")
    if not path:
        return jsonify({"error": "path is required"}), 400
    try:
        return jsonify(delete_path(path))
    except Exception as e:
        return jsonify({"error": str(e)}), 400


@app.route("/api/files/rename", methods=["POST"])
def file_rename():
    data = request.get_json()
    old_path, new_path = data.get("oldPath"), data.get("newPath")
    if not old_path or not new_path:
        return jsonify({"error": "oldPath and newPath are required"}), 400
    try:
        return jsonify(rename_path(old_path, new_path))
    except Exception as e:
        return jsonify({"error": str(e)}), 400


# ─── Export ───────────────────────────────────────────────────────────────────

@app.route("/api/export/docx", methods=["POST"])
def export_docx():
    """
    Generate a richly styled Word (.docx) document from chat messages.

    POST body:
        {
          "title":         str   (optional, default "Vesper Chat Export"),
          "workspaceName": str   (optional),
          "messages": [
            { "role": "user"|"assistant", "content": str,
              "aiId": str (optional), "timestamp": str ISO8601 (optional) }
          ]
        }

    Returns: application/vnd.openxmlformats-officedocument.wordprocessingml.document
    """
    import io, re
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    data          = request.get_json(force=True) or {}
    title         = data.get("title", "Vesper Chat Export")
    workspace     = data.get("workspaceName", "")
    messages_raw  = data.get("messages", [])

    AI_LABELS = {
        "__auto__": "Auto", "pollinations": "Pollinations",
        "llm7": "LLM7",     "chatgpt": "ChatGPT",
    }

    def agent_label(ai_id):
        return AI_LABELS.get(ai_id, ai_id or "Assistant")

    def fmt_ts(ts_str):
        if not ts_str:
            return ""
        try:
            from datetime import datetime, timezone
            dt = datetime.fromisoformat(ts_str.replace("Z", "+00:00"))
            return dt.strftime("%b %d, %Y  %I:%M %p")
        except Exception:
            return ts_str

    def add_shading(paragraph, hex_color="1A1A2E"):
        """Fill paragraph background (code blocks)."""
        pPr = paragraph._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color)
        pPr.append(shd)

    def add_border_bottom(paragraph, hex_color):
        """Add a coloured bottom border to simulate a message header divider."""
        pPr = paragraph._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), hex_color)
        pBdr.append(bottom)
        pPr.append(pBdr)

    def strip_md_inline(text):
        """Strip *bold* / _italic_ / `code` markers for plain paragraph runs."""
        return re.sub(r"(\*\*|__)(.*?)\1", r"\2",
               re.sub(r"(\*|_)(.*?)\1", r"\2",
               re.sub(r"`([^`]+)`", r"\1", text)))

    def add_text_paragraph(doc, line):
        """Render one markdown line as a Word paragraph with inline formatting."""
        h3 = re.match(r"^#{3}\s+(.*)", line)
        h2 = re.match(r"^#{2}\s+(.*)", line)
        h1 = re.match(r"^#{1}\s+(.*)", line)
        if h1:
            p = doc.add_heading(h1.group(1), level=1); return p
        if h2:
            p = doc.add_heading(h2.group(1), level=2); return p
        if h3:
            p = doc.add_heading(h3.group(1), level=3); return p

        # Bullet
        bullet = re.match(r"^[-*+]\s+(.*)", line)
        is_bullet = bool(bullet)
        text_src = bullet.group(1) if bullet else line

        p = doc.add_paragraph(style="List Bullet" if is_bullet else "Normal")
        # Inline tokenise: **bold** | *italic* | `code` | plain
        tokens = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)", text_src)
        for tok in tokens:
            if tok.startswith("**") and tok.endswith("**"):
                run = p.add_run(tok[2:-2]); run.bold = True
            elif tok.startswith("*") and tok.endswith("*"):
                run = p.add_run(tok[1:-1]); run.italic = True
            elif tok.startswith("`") and tok.endswith("`"):
                run = p.add_run(tok[1:-1])
                run.font.name = "Courier New"; run.font.size = Pt(10)
            else:
                p.add_run(tok)
        return p

    # ── Build the document ────────────────────────────────────────────────────
    doc = Document()

    # Page margins (narrow for more content space)
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── Document title ─────────────────────────────────────────────────────────
    title_para = doc.add_heading(title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title_para.runs[0]
    run.font.color.rgb = RGBColor(0x7C, 0x3A, 0xED)  # violet

    meta_parts = []
    if workspace:
        meta_parts.append(f"Workspace: {workspace}")
    meta_parts.append(f"{len(messages_raw)} message{'s' if len(messages_raw) != 1 else ''}")
    from datetime import datetime
    meta_parts.append(f"Exported {datetime.now().strftime('%b %d, %Y  %I:%M %p')}")

    meta_p = doc.add_paragraph("  ·  ".join(meta_parts))
    meta_p.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    meta_p.runs[0].font.size = Pt(9)
    doc.add_paragraph("")  # spacer

    # ── Messages ───────────────────────────────────────────────────────────────
    for msg in messages_raw:
        role      = msg.get("role", "user")
        content   = msg.get("content", "")
        ai_id     = msg.get("aiId", "")
        ts        = fmt_ts(msg.get("timestamp", ""))
        is_user   = role == "user"
        sender    = "You" if is_user else agent_label(ai_id)
        hex_color = "7C3AED" if is_user else "0284C7"

        # Header paragraph: "Sender  ·  timestamp"
        header_p = doc.add_paragraph()
        run_name = header_p.add_run(sender)
        run_name.bold = True
        run_name.font.color.rgb = RGBColor(
            int(hex_color[:2], 16),
            int(hex_color[2:4], 16),
            int(hex_color[4:], 16),
        )
        run_name.font.size = Pt(11)
        if ts:
            run_ts = header_p.add_run(f"   {ts}")
            run_ts.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
            run_ts.font.size = Pt(9)
        add_border_bottom(header_p, hex_color)
        header_p.paragraph_format.space_before = Pt(12)
        header_p.paragraph_format.space_after  = Pt(4)

        # Body — split into text / code segments
        code_pattern = re.compile(r"```(\w*)\n?([\s\S]*?)```", re.DOTALL)
        last = 0
        for m in code_pattern.finditer(content):
            # Text before code block
            text_chunk = content[last:m.start()]
            for line in text_chunk.split("\n"):
                line = line.rstrip()
                if line:
                    add_text_paragraph(doc, line)

            # Code block
            lang = m.group(1).upper() or "CODE"
            code_lines = m.group(2).rstrip().split("\n")
            lang_p = doc.add_paragraph(lang)
            lang_p.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)
            lang_p.runs[0].font.size = Pt(8)
            lang_p.runs[0].font.name = "Courier New"

            for cl in code_lines:
                cp = doc.add_paragraph()
                code_run = cp.add_run(cl if cl else " ")
                code_run.font.name = "Courier New"
                code_run.font.size = Pt(9)
                code_run.font.color.rgb = RGBColor(0xD4, 0xD4, 0xE8)
                add_shading(cp, "1A1A2E")
                cp.paragraph_format.space_before = Pt(0)
                cp.paragraph_format.space_after  = Pt(0)

            last = m.end()

        # Remaining text after last code block
        trailing = content[last:]
        for line in trailing.split("\n"):
            line = line.rstrip()
            if line:
                add_text_paragraph(doc, line)

    # ── Serialise to bytes and return ─────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    slug = re.sub(r"[^a-z0-9-]", "", title.lower().replace(" ", "-"))[:48]
    filename = f"{slug}-{datetime.now().strftime('%Y-%m-%d')}.docx"

    from flask import send_file
    return send_file(
        buf,
        as_attachment=True,
        download_name=filename,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


# ─── Terminal ─────────────────────────────────────────────────────────────────

# RTK-style cumulative token savings tracker
_rtk_stats: dict = {
    "total_original_chars": 0,
    "total_reduced_chars": 0,
    "commands_reduced": 0,
    "commands_total": 0,
}


@app.route("/api/terminal/exec", methods=["POST"])
def terminal_exec():
    from token_reducer import get_stats as _rtk_get_stats
    data = request.get_json()
    command = data.get("command", "").strip()
    if not command:
        return jsonify({"error": "command is required"}), 400
    try:
        result = exec_command(command, cwd=data.get("cwd"), timeout=min(int(data.get("timeout", 60)), 300))
        # Update cumulative RTK savings stats (best-effort)
        try:
            raw_stdout = result.get("stdout", "")
            raw_stderr = result.get("stderr", "")
            stats = _rtk_get_stats(command, raw_stdout, raw_stderr)
            _rtk_stats["commands_total"] += 1
            _rtk_stats["total_original_chars"] += stats["original_chars"]
            _rtk_stats["total_reduced_chars"] += stats["reduced_chars"]
            if stats["savings_pct"] > 0:
                _rtk_stats["commands_reduced"] += 1
        except Exception:
            pass
        return jsonify(result)
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/terminal/savings", methods=["GET"])
def terminal_savings():
    """Return cumulative RTK token savings stats for this session."""
    orig = _rtk_stats["total_original_chars"]
    reduced = _rtk_stats["total_reduced_chars"]
    saved = max(0, orig - reduced)
    pct = int((saved / orig) * 100) if orig > 0 else 0
    return jsonify({
        "originalChars": orig,
        "reducedChars": reduced,
        "savedChars": saved,
        "savingsPct": pct,
        "commandsTotal": _rtk_stats["commands_total"],
        "commandsReduced": _rtk_stats["commands_reduced"],
    })


@app.route("/api/terminal/cwd", methods=["GET"])
def terminal_cwd():
    try:
        return jsonify(get_env_info())
    except Exception as e:
        return jsonify({"error": str(e)}), 500


# ─── Workspaces ───────────────────────────────────────────────────────────────

@app.route("/api/workspaces", methods=["GET"])
def workspaces_list():
    from workspace_manager import list_workspaces
    try:
        return jsonify({"success": True, "workspaces": list_workspaces()})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/workspaces/create", methods=["POST"])
def workspaces_create():
    from workspace_manager import create_workspace
    data = request.get_json() or {}
    name = (data.get("name") or "").strip()
    if not name:
        return jsonify({"error": "'name' is required"}), 400
    try:
        return jsonify(create_workspace(name))
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/workspaces/<workspace_id>/deps", methods=["GET"])
def workspaces_deps(workspace_id):
    from workspace_manager import get_workspace_deps
    try:
        body, status = get_workspace_deps(workspace_id)
        return jsonify(body), status
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/workspaces/<workspace_id>/install", methods=["POST"])
def workspaces_install(workspace_id):
    from workspace_manager import install_dependency
    data = request.get_json() or {}
    package = (data.get("package") or "").strip()
    if not package:
        return jsonify({"error": "'package' is required"}), 400
    version = (data.get("version") or "").strip() or None
    try:
        body, status = install_dependency(workspace_id, package, version)
        return jsonify(body), status
    except Exception as e:
        return jsonify({"error": str(e)}), 500


# ─── Web Scraper ──────────────────────────────────────────────────────────────

@app.route("/api/scraper/scrape", methods=["POST"])
def scraper_scrape():
    data = request.get_json() or {}
    url = (data.get("url") or "").strip()
    if not url:
        return jsonify({"error": "'url' is required"}), 400
    if not url.startswith(("http://", "https://")):
        url = "https://" + url
    selector = data.get("selector") or None
    dynamic = bool(data.get("dynamic", False))
    try:
        from web_scraper import scrape
        result = scrape(url, selector=selector, dynamic=dynamic)
        return jsonify({"result": result, "url": url})
    except Exception as e:
        logger.error(f"Scrape error: {e}", exc_info=True)
        return jsonify({"error": str(e)}), 500


@app.route("/api/scraper/search", methods=["GET", "POST"])
def scraper_search():
    if request.method == "POST":
        data = request.get_json() or {}
        query = (data.get("query") or "").strip()
        num_results = int(data.get("num_results", 8))
    else:
        query = (request.args.get("q") or "").strip()
        num_results = int(request.args.get("n", 8))
    if not query:
        return jsonify({"error": "'query' (or 'q') is required"}), 400
    try:
        from web_scraper import search
        result = search(query, num_results=min(num_results, 10))
        return jsonify({"result": result, "query": query})
    except Exception as e:
        logger.error(f"Search error: {e}", exc_info=True)
        return jsonify({"error": str(e)}), 500


# ─── Multi-Agent Swarm ────────────────────────────────────────────────────────

import multi_agent as _multi_agent
import graph_analyzer as _graph

# ─── Graph analysis (Graphify) ────────────────────────────────────────────────

@app.route("/api/graph/jobs", methods=["GET"])
def graph_jobs():
    return jsonify({"jobs": _graph.list_jobs()})


@app.route("/api/graph/analyze", methods=["POST"])
def graph_analyze():
    data = request.get_json() or {}
    root = (data.get("root") or "").strip()
    if not root:
        root = "/home/runner/workspace"
    ext_filter = data.get("extensions") or None
    job_id = _graph.spawn(root=root, extensions_filter=ext_filter)
    return jsonify({"jobId": job_id, "job": _graph.get_job(job_id)})


@app.route("/api/graph/jobs/<job_id>", methods=["GET"])
def graph_job_status(job_id):
    job = _graph.get_job(job_id)
    if not job:
        return jsonify({"error": "Job not found"}), 404
    return jsonify(job)


@app.route("/api/graph/jobs/<job_id>", methods=["DELETE"])
def graph_job_clear(job_id):
    ok = _graph.clear_job(job_id)
    return jsonify({"cleared": ok})


@app.route("/api/graph/clear-done", methods=["POST"])
def graph_clear_done():
    count = _graph.clear_done_jobs()
    return jsonify({"cleared": count})


@app.route("/api/agents", methods=["GET"])
def agents_list():
    return jsonify({"agents": _multi_agent.list_all()})


@app.route("/api/agents/spawn", methods=["POST"])
def agents_spawn():
    data = request.get_json() or {}
    ai_id = (data.get("aiId") or "").strip()
    task = (data.get("task") or "").strip()
    if not ai_id or not task:
        return jsonify({"error": "aiId and task are required"}), 400
    agent_id = _multi_agent.spawn(
        ai_id=ai_id,
        task=task,
        role=data.get("role", "builder"),
        working_dir=data.get("workingDir") or None,
        max_steps=min(int(data.get("maxSteps", 20)), 50),
        model_id=data.get("modelId") or None,
        label=data.get("label") or None,
    )
    return jsonify({"agentId": agent_id, "status": _multi_agent.get(agent_id)})


@app.route("/api/agents/<agent_id>", methods=["GET"])
def agents_get(agent_id):
    status = _multi_agent.get(agent_id)
    if not status:
        return jsonify({"error": "Agent not found"}), 404
    return jsonify(status)


@app.route("/api/agents/<agent_id>/stop", methods=["POST"])
def agents_stop(agent_id):
    ok = _multi_agent.stop(agent_id)
    return jsonify({"stopped": ok})


@app.route("/api/agents/<agent_id>", methods=["DELETE"])
def agents_clear(agent_id):
    ok = _multi_agent.clear(agent_id)
    return jsonify({"cleared": ok})


@app.route("/api/agents/clear-done", methods=["POST"])
def agents_clear_done():
    count = _multi_agent.clear_all_done()
    return jsonify({"cleared": count})


# ─── Agent ───────────────────────────────────────────────────────────────────

@app.route("/api/agent/run", methods=["POST"])
def agent_run():
    data = request.get_json()
    ai_id = data.get("aiId")
    task = data.get("task", "").strip()
    if not ai_id or not task:
        return jsonify({"error": "aiId and task are required"}), 400

    working_dir = data.get("workingDir")
    max_steps = min(int(data.get("maxSteps", 20)), 50)
    model_id = data.get("modelId") or None
    agent_type = data.get("agentType", "builder")

    if get_agent_status().get("running"):
        return jsonify({"error": "An agent task is already running"}), 409

    def _run():
        try:
            run_agent(ai_id, task, working_dir=working_dir, max_steps=max_steps, model_id=model_id, agent_type=agent_type)
        except Exception as e:
            logger.error("Agent thread error: %s", e)

    threading.Thread(target=_run, daemon=True).start()
    return jsonify({"running": True, "task": task, "steps": [], "result": None})


@app.route("/api/agent/stop", methods=["POST"])
def agent_stop():
    stopped = stop_agent()
    if stopped:
        return jsonify({"success": True, "message": "Stop signal sent to agent"})
    return jsonify({"success": False, "message": "No agent task is currently running"}), 400


@app.route("/api/agent/status", methods=["GET"])
def agent_status():
    try:
        return jsonify(get_agent_status())
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/agent/screenshot/<filename>", methods=["GET"])
def agent_screenshot(filename: str):
    if not filename.endswith(".png") or "/" in filename or ".." in filename:
        return jsonify({"error": "invalid filename"}), 400
    path = get_screenshot_path(filename)
    if not path:
        return jsonify({"error": "screenshot not found"}), 404
    return send_file(str(path), mimetype="image/png")


if __name__ == "__main__":
    port = int(os.environ.get("PYTHON_BACKEND_PORT", 5050))
    logger.info("Starting Universal AI Coding Proxy backend on port %d", port)
    from key_store import migrate_legacy_files, restore_from_kv
    migrate_legacy_files()   # push any leftover local files → KV
    restore_from_kv()        # pull all KV credentials → local files
    app.run(host="0.0.0.0", port=port, debug=False)
